"""
Technical Report Generation Service for ITRC Evaluation System
Generates professional Word documents with evaluation data
"""

import os
import json
from datetime import datetime
from typing import Dict, List, Optional
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

from sqlalchemy.orm import Session
from ..models import TechnicalReport, Evaluation, Application, User, SecurityTarget, STClassSelection, ProductClass
from ..database import get_db

class TechnicalReportGenerator:
    """Professional technical report generator for ITRC evaluation system"""
    
    def __init__(self, db: Session):
        self.db = db
        self.reports_dir = Path("reports")
        self.reports_dir.mkdir(exist_ok=True)
        
    def generate_report_number(self) -> str:
        """Generate unique report number: ITRC-ETR-YYYY-NNNN"""
        year = datetime.now().year
        
        # Count existing reports for the year
        count = self.db.query(TechnicalReport).filter(
            TechnicalReport.report_number.like(f"ITRC-ETR-{year}-%")
        ).count()
        
        return f"ITRC-ETR-{year}-{count + 1:04d}"
    
    def collect_evaluation_data(self, evaluation_id: int) -> Dict:
        """Collect all data needed for report generation"""
        evaluation = self.db.query(Evaluation).filter(Evaluation.id == evaluation_id).first()
        if not evaluation:
            raise ValueError(f"Evaluation {evaluation_id} not found")
        
        application = evaluation.application
        security_target = application.security_target
        # Fix: If security_target is a list, get the first one
        if isinstance(security_target, list):
            if not security_target:
                raise ValueError(f"No security target found for application {application.id}")
            security_target = security_target[0]
        if not security_target:
            raise ValueError(f"No security target found for application {application.id}")
        evaluator = evaluation.evaluator
        applicant = application.applicant
        
        # Get class selections with evaluations
        class_selections = self.db.query(STClassSelection).filter(
            STClassSelection.security_target_id == security_target.id
        ).all()
        
        # Organize class data
        class_data = []
        for selection in class_selections:
            class_info = {
                "class_name_en": selection.product_class.name_en,
                "class_name_fa": selection.product_class.name_fa,
                "class_code": selection.product_class.code,
                "subclass_name_en": selection.product_subclass.name_en if selection.product_subclass else "",
                "subclass_name_fa": selection.product_subclass.name_fa if selection.product_subclass else "",
                "subclass_code": selection.product_subclass.code if selection.product_subclass else "",
                "description": selection.description,
                "justification": selection.justification,
                "test_approach": selection.test_approach,
                "evaluator_notes": selection.evaluator_notes,
                "evaluation_status": selection.evaluation_status,
                "evaluation_score": selection.evaluation_score
            }
            class_data.append(class_info)
        
        return {
            "evaluation": {
                "id": evaluation.id,
                "start_date": evaluation.start_date,
                "end_date": evaluation.end_date,
                "status": evaluation.status,
                "overall_score": evaluation.overall_score,
                "findings": evaluation.findings,
                "recommendations": evaluation.recommendations
            },
            "application": {
                "id": application.id,
                "application_number": application.application_number,
                "product_name": application.product_name,
                "product_version": application.product_version,
                "description": application.description,
                "evaluation_level": application.evaluation_level,
                "company_name": application.company_name,
                "contact_person": application.contact_person,
                "contact_email": application.contact_email,
                "submission_date": application.submission_date
            },
            "evaluator": {
                "id": evaluator.id,
                "full_name": evaluator.full_name,
                "email": evaluator.email,
                "company": evaluator.company
            },
            "applicant": {
                "id": applicant.id,
                "full_name": applicant.full_name,
                "email": applicant.email,
                "company": applicant.company
            },
            "security_target": {
                "id": security_target.id,
                "version": security_target.version,
                "product_description": security_target.product_description,
                "toe_description": security_target.toe_description
            },
            "class_selections": class_data,
            "generation_date": datetime.now(),
            "report_template_version": "1.0"
        }
    
    def create_word_document(self, data: Dict, report_title: str) -> Document:
        """Create Word document with professional formatting"""
        doc = Document()
        
        # Configure document styles
        self._setup_document_styles(doc)
        
        # Title page
        self._add_title_page(doc, data, report_title)
        
        # Table of contents placeholder
        doc.add_page_break()
        self._add_section_heading(doc, "Table of Contents", level=1)
        doc.add_paragraph("(Table of contents will be generated automatically)")
        
        # Executive Summary
        doc.add_page_break()
        self._add_executive_summary(doc, data)
        
        # Product Overview
        self._add_product_overview(doc, data)
        
        # Evaluation Methodology
        self._add_evaluation_methodology(doc, data)
        
        # Detailed Evaluation Results
        self._add_detailed_evaluation_results(doc, data)
        
        # Findings and Recommendations
        self._add_findings_and_recommendations(doc, data)
        
        # Conclusions
        self._add_conclusions(doc, data)
        
        # Appendices
        self._add_appendices(doc, data)
        
        return doc
    
    def _setup_document_styles(self, doc: Document):
        """Setup Persian professional document styles with B Nazanin font and RTL direction"""
        styles = doc.styles
        # Heading styles
        heading1 = styles['Heading 1']
        heading1.font.name = 'B Nazanin'
        heading1._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
        heading1.font.size = Pt(18)
        heading1.font.bold = True
        
        heading2 = styles['Heading 2']
        heading2.font.name = 'B Nazanin'
        heading2._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
        heading2.font.size = Pt(16)
        heading2.font.bold = True
        
        # Normal style
        normal = styles['Normal']
        normal.font.name = 'B Nazanin'
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
        normal.font.size = Pt(14)
    
    def _set_rtl(self, paragraph):
        """Set paragraph direction to RTL"""
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def _add_title_page(self, doc: Document, data: Dict, report_title: str):
        """Add Persian title page with B Nazanin and RTL"""
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_rtl(title_para)
        title_run = title_para.add_run("مرکز تحقیقات مخابرات ایران (ITRC)")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = 'B Nazanin'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
        doc.add_paragraph()
        # Report title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_rtl(title_para)
        title_run = title_para.add_run(report_title)
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.name = 'B Nazanin'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
        doc.add_paragraph()
        # Product info
        product_para = doc.add_paragraph()
        product_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_rtl(product_para)
        product_run = product_para.add_run(f"محصول: {data['application']['product_name']}")
        product_run.font.size = Pt(16)
        product_run.font.name = 'B Nazanin'
        product_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
        if data['application']['product_version']:
            version_para = doc.add_paragraph()
            version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_rtl(version_para)
            version_run = version_para.add_run(f"نسخه: {data['application']['product_version']}")
            version_run.font.size = Pt(14)
            version_run.font.name = 'B Nazanin'
            version_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
        for _ in range(5):
            doc.add_paragraph()
        # Report details table (Persian labels)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self._set_rtl(p)
                    for run in p.runs:
                        run.font.name = 'B Nazanin'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                        run.font.size = Pt(14)
        rows_data = [
            ("شماره درخواست:", data['application']['application_number']),
            ("سطح ارزیابی:", data['application']['evaluation_level']),
            ("ارزیاب:", data['evaluator']['full_name']),
            ("شرکت:", data['application']['company_name']),
            ("دوره ارزیابی:", f"{data['evaluation']['start_date'].strftime('%Y/%m/%d')} تا {data['evaluation']['end_date'].strftime('%Y/%m/%d') if data['evaluation']['end_date'] else 'در حال انجام'}"),
            ("تاریخ گزارش:", data['generation_date'].strftime('%Y/%m/%d'))
        ]
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            for run in table.cell(i, 0).paragraphs[0].runs:
                run.font.bold = True
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
    
    def _add_section_heading(self, doc: Document, heading: str, level: int = 1):
        """Add Persian section heading with B Nazanin and RTL"""
        para = doc.add_heading(heading, level=level)
        self._set_rtl(para)
        for run in para.runs:
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(18 if level == 1 else 16)
    
    def _add_executive_summary(self, doc: Document, data: Dict):
        self._add_section_heading(doc, "خلاصه اجرایی", level=1)
        para = doc.add_paragraph(f"این گزارش نتایج ارزیابی معیارهای مشترک محصول {data['application']['product_name']} نسخه {data['application']['product_version'] or 'نامشخص'} را که توسط مرکز تحقیقات مخابرات ایران انجام شده است، ارائه می‌دهد.")
        self._set_rtl(para)
        for run in para.runs:
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(14)
        para = doc.add_paragraph(f"ارزیابی در سطح {data['application']['evaluation_level']} انجام شده و شامل {len(data['class_selections'])} کلاس عملکردی بوده است. امتیاز کلی ارزیابی: {data['evaluation']['overall_score'] or 'نامشخص'}.")
        self._set_rtl(para)
        for run in para.runs:
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(14)
        if data['evaluation']['overall_score']:
            if data['evaluation']['overall_score'] >= 80:
                result_text = "محصول با موفقیت معیارهای ارزیابی را برآورده کرده است."
            elif data['evaluation']['overall_score'] >= 60:
                result_text = "محصول بیشتر معیارها را با برخی ایرادات جزئی برآورده کرده است."
            else:
                result_text = "محصول نیاز به بهبودهای اساسی برای برآورده کردن معیارها دارد."
            para = doc.add_paragraph(result_text)
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
    
    def _add_product_overview(self, doc: Document, data: Dict):
        self._add_section_heading(doc, "مرور کلی محصول", level=1)
        self._add_section_heading(doc, "شناسنامه محصول", level=2)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self._set_rtl(p)
                    for run in p.runs:
                        run.font.name = 'B Nazanin'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                        run.font.size = Pt(14)
        product_details = [
            ("نام محصول:", data['application']['product_name']),
            ("نسخه محصول:", data['application']['product_version'] or "نامشخص"),
            ("توسعه‌دهنده:", data['application']['company_name']),
            ("سطح ارزیابی:", data['application']['evaluation_level']),
            ("شماره درخواست:", data['application']['application_number'])
        ]
        for i, (label, value) in enumerate(product_details):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            for run in table.cell(i, 0).paragraphs[0].runs:
                run.font.bold = True
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
        self._add_section_heading(doc, "توضیحات محصول", level=2)
        para = doc.add_paragraph(data['application']['description'] or "توضیحی ارائه نشده است.")
        self._set_rtl(para)
        for run in para.runs:
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(14)
        if data['security_target']['toe_description']:
            self._add_section_heading(doc, "هدف ارزیابی (TOE)", level=2)
            para = doc.add_paragraph(data['security_target']['toe_description'])
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
    
    def _add_evaluation_methodology(self, doc: Document, data: Dict):
        self._add_section_heading(doc, "روش‌شناسی ارزیابی", level=1)
        para = doc.add_paragraph("ارزیابی بر اساس معیارهای مشترک فناوری اطلاعات (ISO/IEC 15408) و روش‌شناسی مشترک (ISO/IEC 18045) انجام شده است.")
        self._set_rtl(para)
        for run in para.runs:
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(14)
        self._add_section_heading(doc, "فعالیت‌های ارزیابی", level=2)
        activities = [
            "بررسی و تحلیل مستندات",
            "ارزیابی هدف امنیتی (Security Target)",
            "آزمایش عملکردی",
            "ارزیابی آسیب‌پذیری‌ها",
            "بررسی مدیریت پیکربندی"
        ]
        for activity in activities:
            para = doc.add_paragraph(f"• {activity}")
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
        self._add_section_heading(doc, "دوره ارزیابی", level=2)
        para = doc.add_paragraph(f"ارزیابی از تاریخ {data['evaluation']['start_date'].strftime('%Y/%m/%d')} تا {data['evaluation']['end_date'].strftime('%Y/%m/%d') if data['evaluation']['end_date'] else 'در حال انجام'} انجام شده است.")
        self._set_rtl(para)
        for run in para.runs:
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(14)
    
    def _add_detailed_evaluation_results(self, doc: Document, data: Dict):
        self._add_section_heading(doc, "نتایج تفصیلی ارزیابی", level=1)
        for i, class_selection in enumerate(data['class_selections'], 1):
            self._add_section_heading(doc, f"{i}. {class_selection['class_name_fa']}", level=2)
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._set_rtl(p)
                        for run in p.runs:
                            run.font.name = 'B Nazanin'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                            run.font.size = Pt(14)
            class_details = [
                ("کد کلاس:", class_selection['class_code']),
                ("زیرکلاس:", f"{class_selection['subclass_name_fa']} ({class_selection['subclass_code']})" if class_selection['subclass_code'] else "نامشخص"),
                ("وضعیت:", class_selection['evaluation_status']),
                ("امتیاز:", str(class_selection['evaluation_score']) if class_selection['evaluation_score'] else "نامشخص"),
                ("نام فارسی:", class_selection['class_name_fa'])
            ]
            for j, (label, value) in enumerate(class_details):
                table.cell(j, 0).text = label
                table.cell(j, 1).text = value
                for run in table.cell(j, 0).paragraphs[0].runs:
                    run.font.bold = True
                    run.font.name = 'B Nazanin'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    run.font.size = Pt(14)
            self._add_section_heading(doc, "پیاده‌سازی", level=3)
            para = doc.add_paragraph(class_selection['description'] or "توضیحی ارائه نشده است.")
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
            if class_selection['justification']:
                self._add_section_heading(doc, "توجیه", level=3)
                para = doc.add_paragraph(class_selection['justification'])
                self._set_rtl(para)
                for run in para.runs:
                    run.font.name = 'B Nazanin'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    run.font.size = Pt(14)
            if class_selection['test_approach']:
                self._add_section_heading(doc, "رویکرد آزمون", level=3)
                para = doc.add_paragraph(class_selection['test_approach'])
                self._set_rtl(para)
                for run in para.runs:
                    run.font.name = 'B Nazanin'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    run.font.size = Pt(14)
            if class_selection['evaluator_notes']:
                self._add_section_heading(doc, "ارزیابی ارزیاب", level=3)
                para = doc.add_paragraph(class_selection['evaluator_notes'])
                self._set_rtl(para)
                for run in para.runs:
                    run.font.name = 'B Nazanin'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    run.font.size = Pt(14)
    
    def _add_findings_and_recommendations(self, doc: Document, data: Dict):
        self._add_section_heading(doc, "یافته‌ها و توصیه‌ها", level=1)
        if data['evaluation']['findings']:
            self._add_section_heading(doc, "یافته‌های کلیدی", level=2)
            para = doc.add_paragraph(data['evaluation']['findings'])
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
        if data['evaluation']['recommendations']:
            self._add_section_heading(doc, "توصیه‌ها", level=2)
            para = doc.add_paragraph(data['evaluation']['recommendations'])
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
    
    def _add_conclusions(self, doc: Document, data: Dict):
        self._add_section_heading(doc, "نتیجه‌گیری", level=1)
        overall_score = data['evaluation']['overall_score']
        if overall_score:
            para = doc.add_paragraph(f"بر اساس ارزیابی انجام شده، محصول امتیاز کلی {overall_score}% را کسب کرده است.")
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
            if overall_score >= 80:
                conclusion = "محصول از نظر الزامات امنیتی عملکردی ارزیابی شده، انطباق قوی دارد."
            elif overall_score >= 60:
                conclusion = "محصول با اکثر الزامات انطباق خوبی دارد اما برخی حوزه‌ها نیازمند توجه هستند."
            else:
                conclusion = "محصول نیازمند بهبودهای اساسی برای انطباق رضایت‌بخش است."
            para = doc.add_paragraph(conclusion)
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
        else:
            para = doc.add_paragraph("ارزیابی کامل شده است. لطفاً به یافته‌های تفصیلی مراجعه فرمایید.")
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
    
    def _add_appendices(self, doc: Document, data: Dict):
        self._add_section_heading(doc, "ضمائم", level=1)
        self._add_section_heading(doc, "ضمیمه الف: تیم ارزیابی", level=2)
        para = doc.add_paragraph(f"ارزیاب اصلی: {data['evaluator']['full_name']}")
        self._set_rtl(para)
        for run in para.runs:
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(14)
        para = doc.add_paragraph(f"ایمیل: {data['evaluator']['email']}")
        self._set_rtl(para)
        for run in para.runs:
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(14)
        if data['evaluator']['company']:
            para = doc.add_paragraph(f"سازمان: {data['evaluator']['company']}")
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
        self._add_section_heading(doc, "ضمیمه ب: منابع", level=2)
        references = [
            "ISO/IEC 15408-1:2009، فناوری اطلاعات — تکنیک‌های امنیتی — معیارهای ارزیابی امنیت فناوری اطلاعات — بخش ۱: مقدمه و مدل کلی",
            "ISO/IEC 15408-2:2008، فناوری اطلاعات — تکنیک‌های امنیتی — معیارهای ارزیابی امنیت فناوری اطلاعات — بخش ۲: اجزای عملکردی امنیت",
            "ISO/IEC 15408-3:2008، فناوری اطلاعات — تکنیک‌های امنیتی — معیارهای ارزیابی امنیت فناوری اطلاعات — بخش ۳: اجزای تضمین امنیت",
            "ISO/IEC 18045:2008، فناوری اطلاعات — تکنیک‌های امنیتی — روش‌شناسی ارزیابی امنیت فناوری اطلاعات"
        ]
        for ref in references:
            para = doc.add_paragraph(f"• {ref}")
            self._set_rtl(para)
            for run in para.runs:
                run.font.name = 'B Nazanin'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                run.font.size = Pt(14)
    
    def _serialize_for_json(self, obj):
        """Recursively convert datetime objects in dicts/lists to ISO strings for JSON serialization."""
        if isinstance(obj, dict):
            return {k: self._serialize_for_json(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [self._serialize_for_json(i) for i in obj]
        elif isinstance(obj, datetime):
            return obj.isoformat()
        else:
            return obj
    
    def generate_technical_report(self, evaluation_id: int, generated_by_id: int, 
                                title: Optional[str] = None) -> TechnicalReport:
        """Generate complete technical report"""
        
        # Collect evaluation data
        data = self.collect_evaluation_data(evaluation_id)
        # Serialize datetimes for JSON
        data_serialized = self._serialize_for_json(data)
        
        # Generate report number and title
        report_number = self.generate_report_number()
        if not title:
            title = f"Evaluation Technical Report for {data['application']['product_name']}"
        
        # Create Word document
        doc = self.create_word_document(data, title)
        
        # Save document
        filename = f"{report_number}.docx"
        file_path = self.reports_dir / filename
        doc.save(str(file_path))
        
        # Create database record
        report = TechnicalReport(
            evaluation_id=evaluation_id,
            report_number=report_number,
            title=title,
            generated_by=generated_by_id,
            generated_at=datetime.now(),
            word_file_path=str(file_path),
            file_size=file_path.stat().st_size,
            report_data=data_serialized
        )
        
        self.db.add(report)
        self.db.commit()
        self.db.refresh(report)
        
        return report 